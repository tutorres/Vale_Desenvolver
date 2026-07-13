# -*- coding: utf-8 -*-
"""Fonte única do relatório final.

Este script é a ÚNICA fonte de verdade do documento
Relatorio_Final_Vale_Desenvolver2026.docx. Toda mudança de texto, tabela ou
figura deve ser feita aqui e o .docx é regenerado rodando este script.

Convenções:
- Sem travessões (chr 0x2014). Usar vírgula, parênteses ou dois-pontos.
- Números consumidos de analysis/model_comparison.json e analysis/error_analysis.json.
- Figuras versionadas em figures/ (EDA/SHAP copiadas para lá a partir de data/ e shap_plots/).
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_text(paragraph, text):
    # Clear every w:t, including ones nested inside w:hyperlink (email fields
    # are hyperlink fields in this template, paragraph.runs only sees direct
    # w:r children, so it misses hyperlink-wrapped runs and would otherwise
    # append instead of replace).
    for t in paragraph._p.findall(".//" + qn("w:t")):
        t.text = ""
    for hyperlink in paragraph._p.findall(qn("w:hyperlink")):
        paragraph._p.remove(hyperlink)
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


class Cursor:
    """Insere parágrafos/figuras/tabelas em sequência após um parágrafo âncora,
    mantendo a ordem (cada novo elemento entra logo depois do anterior)."""

    def __init__(self, anchor_paragraph, document):
        self.el = anchor_paragraph._p
        self.parent = anchor_paragraph._parent
        self.doc = document

    def add_para(self, text="", style=None, italic=False, bold=False, center=False):
        new_p = OxmlElement("w:p")
        self.el.addnext(new_p)
        para = Paragraph(new_p, self.parent)
        if style is not None:
            para.style = style
        if text:
            run = para.add_run(text)
            run.italic = italic
            run.bold = bold
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.el = new_p
        return para

    def add_figure(self, path, caption, width_in=6.0):
        pic_p = self.add_para(center=True)
        run = pic_p.add_run()
        run.add_picture(path, width=Inches(width_in))
        self.add_para(caption, italic=True, center=True)

    def add_table(self, rows_data, style="Table Grid"):
        tbl = self.doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        try:
            tbl.style = style
        except Exception:
            pass
        for r, rowvals in enumerate(rows_data):
            for c, val in enumerate(rowvals):
                cell = tbl.cell(r, c)
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(val))
                if r == 0:
                    run.bold = True
        # Reposiciona a tabela (add_table a coloca no fim do corpo) logo após
        # o elemento corrente e avança o cursor para depois dela.
        self.el.addnext(tbl._tbl)
        self.el = tbl._tbl
        return tbl


doc = Document("data/Desenvolver_Template.docx")
p = doc.paragraphs

# --- Participantes ---
set_text(p[2], "")  # Nome do Grupo, trabalho individual
set_text(p[3], "Arthur Torres")
set_text(p[4], "CEFET-MG, Engenharia de Computação")
set_text(p[5], "arthuroliveiratorres@gmail.com")
set_text(p[7], "")
set_text(p[8], "")
set_text(p[9], "")
set_text(p[11], "")
set_text(p[12], "")
set_text(p[13], "")

# --- Resumo ---
resumo = (
    "Este trabalho desenvolve uma solução de análise preditiva para antecipar e apoiar a "
    "prevenção de alertas críticos (Don't Go) em equipamentos de mineração (caminhões e "
    "escavadeiras), a partir de 37.164.054 registros de telemetria coletados entre janeiro e "
    "junho de 2025. Para sustentar essa previsão, a solução foi estruturada como um pipeline "
    "reprodutível em 5 notebooks numerados (limpeza, EDA/validação de regras de negócio, "
    "feature engineering, treinamento de modelo e relatório final), apoiados por 10 módulos "
    "Python testados (93 testes unitários, desenvolvidos com TDD). As features incluem "
    "contagens de alarmes em 5 janelas deslizantes (15min a 4h), tempo desde o último alarme "
    "crítico, tipo de equipamento e estado operacional, todas calculadas sem vazamento "
    "temporal. O modelo é um XGBoost com scale_pos_weight para o desbalanceamento extremo das "
    "classes (1,4% a 2,6% de positivos, conforme o horizonte), treinado e avaliado com split "
    "temporal (treino: janeiro-abril; teste: maio-junho) para 3 horizontes de previsão (1h, "
    "2h, 4h). O modelo escolhido supera com folga baselines de referência (cerca de 4,9x em "
    "F1 sobre o melhor baseline) e foi comparado a uma segunda abordagem (LightGBM), com "
    "matriz de confusão e análise de erros para traduzir o desempenho em impacto operacional."
)
set_text(p[17], resumo)

set_text(p[18], (
    "Palavras-Chave: Telemetria de Mineração; Manutenção Preditiva; XGBoost; LightGBM; "
    "Baseline; Qualidade de Dados; SHAP"
))

# --- Introdução ---
set_text(p[20], (
    "Equipamentos de mineração (caminhões e escavadeiras) geram constantemente alarmes de "
    "telemetria que registram o estado operacional e eventos de risco. Entre eles estão os "
    "alertas \"Don't Go\": situações em que o equipamento acumula um número de alarmes de "
    "determinado tipo dentro de uma janela de tempo, indicando risco operacional que exige "
    "parada ou intervenção. Antecipar esses eventos, em vez de apenas reagir a eles, permite "
    "à operação planejar manutenção, reduzir paradas não programadas e mitigar riscos de "
    "segurança, exemplificando como dados de telemetria, normalmente usados para diagnóstico "
    "retroativo, podem sustentar decisões operacionais proativas."
))
set_text(p[22], (
    "O objetivo central deste trabalho é prever e prevenir a ocorrência de alertas Don't Go: "
    "antecipar, a partir do histórico de alarmes de um equipamento até o instante T, se um "
    "Don't Go ocorrerá nas próximas 1h, 2h ou 4h, para que a operação possa agir antes da "
    "parada, e não apenas registrá-la depois. Como meio para isso, construiu-se um pipeline "
    "robusto e reprodutível que limpa, valida e transforma os dados brutos de telemetria em "
    "um conjunto de features preditivas, sobre o qual se treina um modelo de classificação, "
    "com avaliação honesta de desempenho (F1, precision, recall, AUC) e explicabilidade via "
    "SHAP."
))

# --- Entendimento do Negócio ---
set_text(p[24], (
    "Os dados de telemetria (37.164.054 registros, jan-jun/2025) registram cada alarme "
    "emitido por 35 equipamentos (TAGs), classificados por criticidade (crítico, não crítico, "
    "informacional) e tipo (caminhão ou escavadeira). Paralelamente, os apontamentos (377.907 "
    "registros) descrevem o estado operacional de cada equipamento ao longo do tempo "
    "(Operando, Parado, Hibernando, Manutenção). O sinal Is_Dont_Go, presente na telemetria, "
    "indica quando um equipamento atingiu o critério de parada definido pelas regras de "
    "negócio (136 regras, uma por tipo de alarme, definindo quantidade de ocorrências e "
    "janela de tempo). Esses dados existem porque a operação precisa rastrear, auditar e "
    "justificar paradas de equipamento, tanto por segurança quanto por eficiência de frota. "
    "Antecipar esse sinal, em vez de apenas registrá-lo, é o que transforma esses dados de um "
    "registro histórico em uma ferramenta de decisão: o time de operação pode agir antes da "
    "parada ocorrer, e não apenas explicá-la depois."
))

cur = Cursor(p[24], doc)
cur.add_para((
    "Do ponto de vista de negócio, o sucesso da solução é medido pela sua capacidade de "
    "antecipar Don't Go com antecedência útil para a operação: a métrica-alvo primária é a "
    "fração de Don't Go reais sinalizados antes de ocorrerem (recall), condicionada a um "
    "volume de alertas que a equipe consiga triar (precision operacionalmente aceitável). Um "
    "cenário concreto de aplicação: um alerta no painel do dispatcher indicando, com até 4 "
    "horas de antecedência, que determinado equipamento tem risco elevado de Don't Go, "
    "permitindo reprogramar a frota, antecipar inspeção ou manutenção e evitar a parada não "
    "planejada no meio do ciclo produtivo. Nesse enquadramento, cada Don't Go antecipado é "
    "uma parada potencialmente convertida de não planejada em planejada."
), style="List Bullet")
cur.add_figure(
    "figures/eda_dontgo_timeline.png",
    "Figura 1: Série temporal de ocorrências de Don't Go ao longo do período (jan-jun/2025).",
)

# --- Metodologia ---
set_text(p[26], (
    "O pipeline é dividido em 5 notebooks numerados e um conjunto de 10 módulos Python "
    "testados (src/loader.py, cleaner.py, validator.py, business_rules.py, features.py, "
    "model.py, utils.py, além de baselines.py, lgbm_model.py e error_analysis.py para a "
    "comparação de modelos e a análise de erros, 93 testes unitários, desenvolvidos com "
    "TDD). A arquitetura segue: (1) "
    "carregamento e limpeza dos 6 arquivos mensais de telemetria e dos apontamentos; (2) EDA "
    "e validação das regras de negócio, reproduzindo o sinal Is_Dont_Go via janela deslizante "
    "e comparando com o sinal registrado; (3) feature engineering: contagens de alarmes em 5 "
    "janelas (15min, 30min, 1h, 2h, 4h), tempo desde o último alarme crítico, tipo de "
    "equipamento e estado operacional, calculadas com pandas .rolling() e np.searchsorted "
    "(O(n log n) por equipamento, vetorizado, essencial para os 37M registros); (4) "
    "treinamento de um XGBoost com scale_pos_weight para o desbalanceamento de classes, split "
    "temporal (treino jan-abr, teste mai-jun, nunca aleatório, para não vazar informação do "
    "futuro) e avaliação em 3 horizontes; (5) explicabilidade via SHAP e consolidação do "
    "relatório final."
))
set_text(p[27], (
    "A principal dificuldade técnica foi escalar o pipeline das fixtures sintéticas usadas nos "
    "testes (5 a 15 linhas) para os 37M registros reais, em uma máquina com 16GB de RAM. "
    "Colunas de string carregadas como object (em vez de category) projetavam ~30GB de uso de "
    "memória; corrigido convertendo para category no carregamento. Um problema mais sutil "
    "surgiu na concatenação dos 6 meses: o pandas reverte category para object quando o "
    "conjunto de categorias difere entre os meses (TAGs e operadores que só aparecem em parte "
    "do ano), corrigido unificando o CategoricalDtype antes do concat. O mesmo padrão de "
    "memória se repetiu na feature current_state (apenas 5 valores possíveis, mas calculada "
    "como object) e na cópia de colunas não utilizadas dentro de build_features. Cada um "
    "desses problemas foi reproduzido com um teste antes da correção (TDD), e a suíte completa "
    "(93 testes) permanece verde. Adicionalmente, o dataset real revelou um quarto erro de "
    "qualidade não documentado (string 'NULL' também em Valor, 237.443 registros) e expôs uma "
    "incompatibilidade de versão entre shap e xgboost na etapa de explicabilidade, "
    "documentados em spec/DECISIONS.md (D008)."
))
cur = Cursor(p[27], doc)
cur.add_figure(
    "figures/feature_correlation.png",
    "Figura 2: Matriz de correlação entre as features preditivas engenheiradas.",
)

# --- Resultados e Discussões ---
set_text(p[29], (
    "O diagnóstico de qualidade identificou e corrigiu 4 problemas nos 37.164.054 registros "
    "de telemetria: 11 registros com corrupção de encoding em Criticidade, 36.104.611 "
    "registros com string literal 'NULL' em Classe, 237.443 registros com a mesma string "
    "'NULL' em Valor, e 821.849 registros com vírgula como separador decimal em Valor. A "
    "reprodução do sinal Is_Dont_Go via janela deslizante, validada em uma amostra de 50 mil "
    "registros contra as 136 regras de negócio, obteve precision de 0,12 e recall de 0,59, "
    "indicando que as regras documentadas capturam a maioria dos casos reais, mas com taxa "
    "relevante de falsos positivos."
))

cur = Cursor(p[29], doc)

# Desempenho por horizonte (XGBoost)
cur.add_para((
    "O modelo XGBoost foi treinado e avaliado para os 3 horizontes propostos, com split "
    "temporal (23.273.520 registros de treino, 13.890.534 de teste). O horizonte de 4 horas "
    "apresentou o melhor F1-score (0,186 no threshold 0,5; 0,207 no threshold 0,7), superando "
    "1h (F1=0,114) e 2h (F1=0,143), mesmo com AUC mais baixo (0,767 vs. 0,883 em 1h), o que "
    "indica uma fronteira de decisão mais bem calibrada para a métrica de negócio (F1), ainda "
    "que a capacidade geral de ranqueamento (AUC) seja menor. Em todos os horizontes, o "
    "recall é consideravelmente maior que a precision (ex: recall=0,577 vs. precision=0,111 "
    "em 4h/threshold 0,5), esperado dado o desbalanceamento extremo das classes e a escolha "
    "de scale_pos_weight, que privilegia não perder eventos reais de Don't Go ao custo de "
    "mais falsos positivos. Aumentar o threshold para 0,7 melhora a precision (de 0,111 para "
    "0,126) com perda mínima de recall (de 0,577 para 0,574), sugerindo que esse ajuste é "
    "preferível ao padrão de 0,5 em produção."
), style="Normal")

# Baseline + LightGBM + tabela comparativa
cur.add_para((
    "Para situar esse desempenho, o modelo foi comparado a baselines de referência e a uma "
    "segunda abordagem (LightGBM), todos avaliados no mesmo split temporal e nas mesmas 8 "
    "features (label_4h). Os baselines vão do trivial (classificador que sempre prevê a "
    "classe majoritária, F1=0,000; e um estratificado, F1=0,022) a uma heurística de domínio "
    "(prever Don't Go quando há alarme crítico na última hora, F1=0,038). A base rate de "
    "positivos é de apenas 1,6%. Contra esse pano de fundo, o XGBoost alcança F1=0,186 e "
    "AUC-PR=0,205, superando o melhor baseline em cerca de 4,9x em F1 e cerca de 12,5x em "
    "AUC-PR sobre a base rate: o ganho é substancial e confirma que o modelo aprende sinal "
    "real, não apenas a proporção das classes. O LightGBM empata na prática (F1=0,185, "
    "AUC-PR=0,197), com AUC-ROC e recall marginalmente melhores (0,776 e 0,579), o que "
    "reforça a robustez do resultado por concordância entre duas famílias de gradient "
    "boosting. O XGBoost foi mantido como modelo escolhido por liderar em F1 e AUC-PR, as "
    "métricas mais informativas sob desbalanceamento extremo."
), style="Normal")

cur.add_para(
    "Tabela 1: Comparação de desempenho no split temporal de teste (label_4h, threshold 0,5).",
    italic=True, center=True,
)
cur.add_table([
    ["Modelo", "Precision", "Recall", "F1", "AUC-ROC", "AUC-PR"],
    ["Dummy (mais frequente)", "0,000", "0,000", "0,000", "0,500", "0,016"],
    ["Dummy (estratificado)", "0,017", "0,032", "0,022", "0,500", "0,016"],
    ["Heurístico (critical_1h>0)", "0,021", "0,210", "0,038", "0,529", "0,072"],
    ["XGBoost (escolhido)", "0,111", "0,577", "0,186", "0,767", "0,205"],
    ["LightGBM", "0,110", "0,579", "0,185", "0,776", "0,197"],
])

cur.add_figure(
    "figures/model_comparison.png",
    "Figura 3: Comparação de F1 e recall entre baselines e os modelos (label_4h).",
)
cur.add_figure(
    "figures/roc_pr_curves.png",
    "Figura 4: Curvas ROC e Precision-Recall dos modelos avaliados (label_4h).",
)

# Matriz de confusão + análise de erros
cur.add_para((
    "A matriz de confusão do XGBoost (label_4h, threshold 0,7, teste de maio-junho, "
    "N=13.890.534) detalha o comportamento operacional do modelo: 130.705 verdadeiros "
    "positivos (TP), 907.848 falsos positivos (FP), 96.918 falsos negativos (FN) e 12.755.063 "
    "verdadeiros negativos (TN), para 227.623 Don't Go reais no período (recall 57,4%, "
    "precision 12,6%). Em termos de impacto, cada FN é um Don't Go não antecipado, ou seja, "
    "uma parada não planejada que passou sem alerta, enquanto cada FP é um alerta sem Don't Go "
    "correspondente, que gera uma inspeção potencialmente desnecessária. A análise dos falsos "
    "negativos revela um insight não óbvio: 98,0% deles (94.932 de 96.918) ocorrem em "
    "escavadeiras e apenas 2,0% em caminhões, e 95,7% acontecem com o equipamento no estado "
    "Operando; além disso, 99,9% caem na faixa de alarm_count_4h maior ou igual a 21, "
    "indicando que o modelo falha justamente quando há muitos alarmes acumulados em "
    "escavadeiras. Há ainda um drift temporal marcante: o recall é de apenas 42,3% em maio, "
    "mas sobe para 92,2% em junho, concentrando a maioria dos FN no primeiro mês de teste. "
    "Como limitação honesta, a precision baixa (12,6%) significa que a maior parte dos alertas "
    "é falso positivo, o que exige uma etapa de triagem, e a lacuna de recall em escavadeiras "
    "somada ao drift de maio indica que o modelo ainda não generaliza de forma estável entre "
    "tipos de equipamento e períodos."
), style="Normal")
cur.add_figure(
    "figures/confusion_matrix_4h.png",
    "Figura 5: Matriz de confusão do XGBoost (label_4h, threshold 0,7) com impacto operacional.",
)

# Impacto de negócio quantificado
cur.add_para((
    "Traduzindo o desempenho em impacto de negócio: com recall de 57,4% no threshold 0,7, o "
    "modelo teria antecipado cerca de 130.700 dos 227.623 Don't Go do bimestre de teste "
    "(maio-junho), sinalizando-os com até 4 horas de antecedência. O valor em horas de parada "
    "evitada depende de uma premissa operacional que deve ser calibrada com a operação: se "
    "cada Don't Go antecipado permite converter em planejada uma parada não planejada que, na "
    "média, custaria H horas de indisponibilidade, então o modelo evitaria da ordem de "
    "130.700 x H horas de parada não planejada no período. A título ilustrativo, adotando a "
    "premissa conservadora de H = 1 hora por evento, isso corresponde a aproximadamente "
    "130.700 horas de parada não planejada mitigadas no bimestre; o número escala "
    "linearmente com H e deve ser substituído pelo tempo médio real de parada assim que "
    "disponível. Essa estimativa não considera o custo dos falsos positivos (inspeções "
    "desnecessárias), que deve ser ponderado contra o benefício ao definir o threshold de "
    "operação."
), style="Normal")

# SHAP
cur.add_para((
    "A análise SHAP para o horizonte de 4h revelou que tipo_caminhao (caminhão vs. "
    "escavadeira) é, isoladamente, a feature mais influente, mais que o dobro do impacto de "
    "time_since_last_critical, a segunda colocada, e muito acima das contagens de alarme em "
    "qualquer janela. Esse resultado é uma limitação relevante a ser discutida com "
    "transparência: sugere que o modelo aprendeu majoritariamente uma diferença de risco-base "
    "entre os dois tipos de equipamento, em vez de um padrão dinâmico fino de comportamento de "
    "alarmes antes de um Don't Go, o que é coerente com a concentração de falsos negativos em "
    "escavadeiras observada na matriz de confusão. As features de contagem de alarme "
    "contribuem, mas com impacto bem menor, alarm_count_4h é a terceira mais relevante, "
    "seguida por uma queda abrupta nas janelas mais curtas (30m, 1h, 2h, 15m), coerente com o "
    "horizonte de 4h ser o melhor: alarmes acumulados em uma janela mais longa carregam mais "
    "sinal preditivo para esse horizonte do que rajadas de curto prazo."
), style="Normal")
cur.add_figure(
    "figures/shap_summary_label_4h.png",
    "Figura 6: Resumo SHAP das features para o horizonte de 4h (impacto por feature).",
)

# --- Conclusão e Trabalhos Futuros ---
set_text(p[31], (
    "Este trabalho demonstra que é possível antecipar e, com isso, apoiar a prevenção de "
    "ocorrências de Don't Go com até 4 horas de antecedência, usando exclusivamente features "
    "derivadas do histórico de alarmes de telemetria, apoiado por um pipeline auditável e "
    "testado (93 testes, TDD) que escala dos dados sintéticos de teste até os 37 milhões de "
    "registros reais. O modelo escolhido supera com folga os baselines de referência (cerca "
    "de 4,9x em F1 sobre o melhor baseline) e, com recall de 57%, teria antecipado cerca de "
    "130.700 dos 227.623 Don't Go do bimestre de teste. As métricas absolutas (F1 entre 0,11 "
    "e 0,21) são modestas, mas refletem honestamente um problema com desbalanceamento extremo "
    "(abaixo de 3% de positivos), uma feature dominante (tipo de equipamento) e uma lacuna de "
    "recall em escavadeiras, que apontam espaço real para melhoria nas features dinâmicas "
    "antes de uma eventual adoção em produção."
))
set_text(p[32], (
    "Investigar por que tipo_caminhao domina o SHAP: treinar modelos separados por tipo de "
    "equipamento, ou adicionar interações explícitas, para isolar o sinal dinâmico de alarmes "
    "do risco-base por tipo."
))
b2 = insert_paragraph_after(p[32], (
    "Tratar a lacuna de recall em escavadeiras, onde se concentram 98% dos falsos negativos, "
    "com calibração ou modelo específico para esse tipo de equipamento."
), style="List Bullet")
b3 = insert_paragraph_after(b2, (
    "Mitigar o drift temporal observado entre maio e junho (recall de 42% para 92%) com "
    "re-treinamento periódico e monitoramento contínuo da distribuição das features."
), style="List Bullet")
b4 = insert_paragraph_after(b3, (
    "Adicionar features de turno, frequência de manutenções recentes e localidade."
), style="List Bullet")
insert_paragraph_after(b4, (
    "Construir um dashboard (ex: Streamlit) para consulta de risco por equipamento em tempo "
    "real, materializando o alerta do dispatcher descrito no entendimento do negócio."
), style="List Bullet")

doc.save("Relatorio_Final_Vale_Desenvolver2026.docx")
print("Saved Relatorio_Final_Vale_Desenvolver2026.docx")
